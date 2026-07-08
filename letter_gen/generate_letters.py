#!/usr/bin/env python3
"""
Generalized letter generator.

Reads:
  - config.json   -> maps each document type to a template, a location in
                      data.json, and a placeholder -> field mapping
  - data.json      -> the OCR/master JSON, can contain many document sections

For every document type in config.json, this script:
  1. Locates its data inside data.json (supports nested paths and lists)
  2. Skips it gracefully if the template file doesn't exist yet, or if the
     required fields are missing/empty (e.g. 12th marksheet with no data)
  3. Fills the template's {{placeholders}} and saves a separate output file
     per record (so 3 employment certificates -> 3 separate letters)

Usage:
    python3 generate_letters.py
"""
import json
import logging
import re
from pathlib import Path
from datetime import date
from docx import Document

logging.basicConfig(level=logging.INFO, format="%(message)s")
log = logging.getLogger("letter_gen")


# --------------------------------------------------------------------------- #
# Data resolution helpers
# --------------------------------------------------------------------------- #

def get_nested(data, path):
    """Resolve a dotted path like 'employment_certificate.certificates'."""
    cur = data
    for part in path.split("."):
        if not isinstance(cur, dict) or part not in cur:
            return None
        cur = cur[part]
    return cur


def resolve_field(item: dict, spec: str):
    """
    Resolve one field spec against a record dict.
      "field_name"          -> item["field_name"]
      "field_name|upper"    -> item["field_name"].upper()
      "__today__"           -> today's date, formatted
      "__const__:text"      -> literal constant text
      "__address__:key"     -> join non-null parts of item["key"] (a sub-dict)
    """
    if spec == "__today__":
        today = date.today()
        return f"{today.day} {today.strftime('%B %Y')}"

    if spec.startswith("__const__:"):
        return spec.split(":", 1)[1]

    if spec.startswith("__address__:"):
        sub = item.get(spec.split(":", 1)[1]) or {}
        if not isinstance(sub, dict):
            return str(sub) if sub else ""
        parts = [
            sub.get("address_line"),
            sub.get("city_or_village"),
            sub.get("district"),
            sub.get("state"),
            sub.get("pincode"),
            sub.get("country"),
        ]
        return ", ".join(str(p) for p in parts if p)

    field, _, transform = spec.partition("|")
    value = item.get(field)
    value = "" if value is None else str(value)
    if transform == "upper":
        value = value.upper()
    return value


def has_required(item: dict, required: list) -> bool:
    return all(item.get(f) not in (None, "", []) for f in required)


# --------------------------------------------------------------------------- #
# docx traversal
# --------------------------------------------------------------------------- #

def iter_paragraphs(container):
    """Yield every paragraph in a body/cell, including nested tables."""
    if hasattr(container, "paragraphs"):
        for paragraph in container.paragraphs:
            yield paragraph

    if hasattr(container, "tables"):
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)


def iter_all_paragraphs(document: Document):
    """Yield paragraphs from the body AND from every section's headers/footers."""
    yield from iter_paragraphs(document)

    for section in document.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if part is not None:
                yield from iter_paragraphs(part)


# --------------------------------------------------------------------------- #
# Placeholder replacement — formatting-safe
# --------------------------------------------------------------------------- #

def _build_placeholder_pattern(placeholders: dict):
    """One regex that matches any of the {{placeholder}} keys, tolerant of
    stray whitespace inside the braces (Word/OCR sometimes inserts it)."""
    inner_alternatives = [re.escape(ph[2:-2]) for ph in placeholders]
    return re.compile(r"\{\{\s*(" + "|".join(inner_alternatives) + r")\s*\}\}")


def replace_placeholders_in_paragraph(paragraph, placeholders: dict, pattern: re.Pattern) -> set:
    """
    Replace {{placeholder}} occurrences in a paragraph WITHOUT touching runs
    that don't contain a placeholder, so unrelated formatting (bold, font,
    color, size) elsewhere in the paragraph is left completely alone.

    - If a placeholder sits entirely inside one run, only that run's text is
      edited; its formatting is preserved and nothing else is touched.
    - If a placeholder happens to be split across multiple runs (Word does
      this sometimes, e.g. after spellcheck/autocorrect), only the runs that
      the placeholder actually spans are merged — using the first run's
      formatting for the inserted value — everything outside that span is
      untouched.
    """
    runs = paragraph.runs
    if not runs:
        return set()

    run_texts = [r.text or "" for r in runs]
    full_text = "".join(run_texts)
    matches = list(pattern.finditer(full_text))
    if not matches:
        return set()

    # char offset -> run index
    boundaries = []
    offset = 0
    for i, t in enumerate(run_texts):
        boundaries.append((offset, offset + len(t)))
        offset += len(t)

    def run_for_pos(pos):
        for i, (start, end) in enumerate(boundaries):
            if start <= pos < end:
                return i
        return len(runs) - 1

    value_by_key = {ph[2:-2].strip(): v for ph, v in placeholders.items()}

    # Apply replacements back-to-front so earlier offsets stay valid.
    for m in reversed(matches):
        key = m.group(1).strip()
        value = value_by_key.get(key)
        if value is None:
            continue

        start, end = m.span()
        start_run = run_for_pos(start)
        end_run = run_for_pos(end - 1)

        if start_run == end_run:
            local_start = start - boundaries[start_run][0]
            local_end = end - boundaries[start_run][0]
            r = runs[start_run]
            r.text = r.text[:local_start] + value + r.text[local_end:]
        else:
            # Placeholder spans multiple runs: merge only that span.
            local_start = start - boundaries[start_run][0]
            local_end = end - boundaries[end_run][0]
            prefix = runs[start_run].text[:local_start]
            suffix = runs[end_run].text[local_end:]
            runs[start_run].text = prefix + value + suffix
            for i in range(start_run + 1, end_run + 1):
                runs[i].text = ""

    remaining_text = "".join(r.text or "" for r in paragraph.runs)
    return {m.group(0) for m in pattern.finditer(remaining_text)}


def fill_docx(template_path: Path, placeholders: dict, output_path: Path) -> Path:
    document = Document(str(template_path))
    pattern = _build_placeholder_pattern(placeholders)

    remaining_placeholders = set()
    for paragraph in iter_all_paragraphs(document):
        remaining_placeholders.update(
            replace_placeholders_in_paragraph(paragraph, placeholders, pattern)
        )

    if remaining_placeholders:
        log.warning("    Warning: unresolved placeholders %s", remaining_placeholders)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        document.save(str(output_path))
        return output_path
    except PermissionError:
        fallback_path = output_path.with_name(f"{output_path.stem}_generated{output_path.suffix}")
        counter = 1
        while fallback_path.exists():
            fallback_path = output_path.with_name(f"{output_path.stem}_generated_{counter}{output_path.suffix}")
            counter += 1
        document.save(str(fallback_path))
        log.warning("    Warning: output file was open, saved as %s", fallback_path.name)
        return fallback_path


def safe_filename(pattern: str, item: dict) -> str:
    class SafeDict(dict):
        def __missing__(self, key):
            return "unknown"
    name = pattern.format_map(SafeDict(item))
    name = re.sub(r'[\\/:*?"<>|]', "_", name)
    return name


# --------------------------------------------------------------------------- #
# Orchestration
# --------------------------------------------------------------------------- #

def generate_all(data_path: str, config_path: str, outputs_dir: str):
    with open(data_path, encoding="utf-8") as f:
        data = json.load(f)
    with open(config_path, encoding="utf-8") as f:
        config = json.load(f)

    generate_all_from_data(data, config, outputs_dir, Path(config_path).parent)


def generate_all_from_data(data: dict, config: dict, outputs_dir: str, base_dir: Path):
    outputs_dir = Path(outputs_dir)
    outputs_dir.mkdir(parents=True, exist_ok=True)

    for doc_type, conf in config.items():
        log.info("\n[%s]", doc_type)
        template_path = Path(conf["template"])
        if not template_path.is_absolute():
            template_path = base_dir / template_path
        if not template_path.exists():
            log.info("    Skipped - template not found: %s", template_path)
            continue

        source = get_nested(data, conf["source"])
        if not source:
            log.info("    Skipped - no data found at '%s' in JSON", conf["source"])
            continue

        items = source if conf.get("is_list") else [source]

        generated = 0
        for idx, item in enumerate(items):
            if not isinstance(item, dict) or not has_required(item, conf["required"]):
                missing = [f for f in conf["required"] if not isinstance(item, dict) or item.get(f) in (None, "", [])]
                log.info("    Skipped record #%d - missing required fields %s", idx, missing)
                continue

            placeholders = {ph: resolve_field(item, spec) for ph, spec in conf["fields"].items()}
            out_name = safe_filename(conf["output_pattern"], item)
            out_path = outputs_dir / out_name

            try:
                saved_path = fill_docx(template_path, placeholders, out_path)
                log.info("    Created: %s", saved_path)
                generated += 1
            except Exception as exc:
                log.error("    Failed record #%d (%s): %s", idx, out_name, exc)

        if generated == 0:
            log.info("    No letters generated for this document type.")


if __name__ == "__main__":
    data_path = r"E:\AiProff-BCP\BGV Backend\BGV\letter_gen\data\data_ocr.json"
    config_path = r"E:\AiProff-BCP\BGV Backend\BGV\letter_gen\config.json"
    outputs_dir = r"E:\AiProff-BCP\BGV Backend\BGV\letter_gen\output"
    generate_all(data_path, config_path, outputs_dir)